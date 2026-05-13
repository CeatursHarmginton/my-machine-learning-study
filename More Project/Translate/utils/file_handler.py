"""
File handling utilities for reading/writing txt and docx files.
"""
import os
import chardet
from pathlib import Path


def read_file(filepath: str) -> str:
    """
    Read text from a file, supporting .txt and .docx formats.
    Automatically detects encoding for txt files.

    Args:
        filepath: Path to the file.

    Returns:
        File contents as string.

    Raises:
        ValueError: If file format is not supported.
        FileNotFoundError: If file doesn't exist.
    """
    path = Path(filepath)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {filepath}")

    ext = path.suffix.lower()

    if ext == '.txt':
        return _read_txt(path)
    elif ext in ('.docx', '.doc'):
        return _read_docx(path)
    else:
        raise ValueError(f"Unsupported file format: {ext}. Supported: .txt, .docx")


def _read_txt(path: Path) -> str:
    """Read txt file with automatic encoding detection."""
    # Read raw bytes for encoding detection
    raw = path.read_bytes()

    # Detect encoding
    detected = chardet.detect(raw)
    encoding = detected.get('encoding', 'utf-8') or 'utf-8'

    try:
        return raw.decode(encoding)
    except (UnicodeDecodeError, LookupError):
        # Fallback encodings common for CJK texts
        for fallback in ['utf-8', 'utf-8-sig', 'gb18030', 'gbk', 'big5', 'euc-kr', 'shift_jis']:
            try:
                return raw.decode(fallback)
            except (UnicodeDecodeError, LookupError):
                continue
        raise UnicodeDecodeError(f"Could not decode file with any known encoding: {path}")


def _read_docx(path: Path) -> str:
    """Read docx file and extract text content."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx is required for .docx support. "
            "Install it with: pip install python-docx"
        )

    doc = Document(str(path))
    paragraphs = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    return "\n\n".join(paragraphs)


def write_file(filepath: str, content: str, format: str = "txt") -> str:
    """
    Write translation output to file.

    Args:
        filepath: Output file path.
        content: Text content to write.
        format: Output format ("txt" or "docx").

    Returns:
        Path to the written file.
    """
    path = Path(filepath)
    path.parent.mkdir(parents=True, exist_ok=True)

    if format == "txt":
        path.write_text(content, encoding='utf-8')
    elif format == "docx":
        _write_docx(path, content)
    else:
        raise ValueError(f"Unsupported output format: {format}")

    return str(path)


def _write_docx(path: Path, content: str):
    """Write content to a docx file."""
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx is required for .docx output.")

    doc = Document()

    # Split by double newline for paragraphs
    paragraphs = content.split('\n\n')
    for para_text in paragraphs:
        if para_text.strip():
            doc.add_paragraph(para_text.strip())

    doc.save(str(path))


def get_file_info(filepath: str) -> dict:
    """Get file metadata."""
    path = Path(filepath)
    if not path.exists():
        return {"error": "File not found"}

    return {
        "name": path.name,
        "size_bytes": path.stat().st_size,
        "size_readable": _format_size(path.stat().st_size),
        "format": path.suffix.lower(),
        "path": str(path),
    }


def _format_size(size_bytes: int) -> str:
    """Format byte size to human readable."""
    for unit in ['B', 'KB', 'MB', 'GB']:
        if size_bytes < 1024:
            return f"{size_bytes:.1f} {unit}"
        size_bytes /= 1024
    return f"{size_bytes:.1f} TB"
